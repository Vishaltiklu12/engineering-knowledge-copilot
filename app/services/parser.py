from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.exceptions import ValidationAppError


@dataclass
class ParsedPage:
    page_number: int | None
    text: str


@dataclass
class ParsedDocument:
    pages: list[ParsedPage]


class DocumentParser:
    def parse(self, file_path: Path, mime_type: str) -> ParsedDocument:
        suffix = file_path.suffix.lower()

        if suffix in {".txt", ".md"} or mime_type.startswith("text/"):
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return ParsedDocument(pages=[ParsedPage(page_number=1, text=text)])

        if suffix == ".pdf":
            reader = PdfReader(str(file_path))
            pages = [
                ParsedPage(page_number=index + 1, text=(page.extract_text() or "").strip())
                for index, page in enumerate(reader.pages)
            ]
            return ParsedDocument(pages=pages)

        if suffix == ".docx":
            doc = DocxDocument(str(file_path))
            paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
            return ParsedDocument(pages=[ParsedPage(page_number=1, text="\n".join(paragraphs))])

        raise ValidationAppError(
            "Unsupported document type.",
            details={"mime_type": mime_type, "suffix": suffix},
        )
